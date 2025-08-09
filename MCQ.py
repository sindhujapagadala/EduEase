import streamlit as st
import google.generativeai as genai
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import os
from dotenv import load_dotenv

# Load API keys from .env
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

# Configure Gemini API
genai.configure(api_key=GEMINI_API_KEY)

# Function to get MCQ questions from Gemini
def generate_mcq_questions(topic, difficulty, num_questions):
    prompt = f"""
    Generate a multiple-choice quiz with the following specifications:
    - Topic: "{topic}"
    - Difficulty level: "{difficulty}"
    - Number of questions: {num_questions}
    - Each question should have 4 options labeled a, b, c, and d.
    - Each option should be brief (2-3 words).
    - Clearly specify the correct answer for each question.

    Ensure:
    1. Questions and options are clear and concise.
    2. Relevant to the specified topic.
    3. Appropriately challenging for the specified difficulty level.
    4. No code-related questions.

    Example format:
    Q1: What is the capital of France?
    a. Berlin
    b. Madrid
    c. Paris
    d. Rome
    Answer: c

    Quiz:
    """

    model = genai.GenerativeModel("gemini-2.0-flash")
    response = model.generate_content(prompt)
    return response.text.strip()

# Function to format quiz into a structured list
def format_quiz(quiz):
    lines = quiz.split("\n")
    formatted_quiz = []
    current_question = []

    for line in lines:
        if re.match(r"^Q\d+: ", line):
            if current_question:
                formatted_quiz.append(current_question)
            current_question = [line]
        elif re.match(r"^[a-d]\. ", line):
            current_question.append(line)
        elif re.match(r"^Answer: ", line):
            current_question.append(line)
    if current_question:
        formatted_quiz.append(current_question)
    return formatted_quiz

# Function to generate DOCX
def generate_docx(quiz, heading1, heading2):
    doc = Document()
    
    heading1_paragraph = doc.add_heading(heading1, level=0)
    heading1_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    heading2_paragraph = doc.add_heading(heading2, level=2)
    heading2_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("Name:")
    doc.add_paragraph("Roll number:")
    doc.add_paragraph("Class:")
    doc.add_paragraph("Section:")
    doc.add_paragraph("")  # Add space

    for question in quiz:
        for line in question:
            if not re.match(r"^Answer: ", line):
                doc.add_paragraph(line)
        doc.add_paragraph("")

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# Streamlit App
def MCQ():
    st.title("MCQ Quiz Generator")

    institute_name = st.text_input("Enter the Name of Institute:")
    quiz_title = st.text_input("Enter the quiz title:")
    topic = st.text_input("Enter the topic:")
    difficulty = st.selectbox("Select difficulty:", ["Beginner", "Intermediate", "Expert"])
    num_questions = st.number_input("Number of questions:", min_value=1, max_value=20, value=5)

    if st.button("Generate Quiz"):
        if topic:
            with st.spinner("Generating questions using Gemini..."):
                quiz = generate_mcq_questions(topic, difficulty, num_questions)
                formatted_quiz = format_quiz(quiz)
                st.subheader("Generated Quiz:")
                if institute_name:
                    st.write(f"**{institute_name}**")
                if quiz_title:
                    st.write(f"**{quiz_title}**")
                
                st.session_state['quiz'] = formatted_quiz
                st.session_state['docx_content'] = generate_docx(formatted_quiz, institute_name, quiz_title)
        else:
            st.error("Please enter a topic.")

    if 'quiz' in st.session_state:
        for question in st.session_state['quiz']:
            st.write(question[0])
            for line in question[1:]:
                if "Answer: " in line:
                    st.write(f"**Correct answer:** {line.split(': ')[1]}")
                else:
                    st.write(line)

        if 'docx_content' in st.session_state:
            st.download_button(
                label="Download Quiz as DOCX",
                data=st.session_state['docx_content'],
                file_name=f"{topic} quiz.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    MCQ()
