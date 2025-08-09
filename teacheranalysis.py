import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import google.generativeai as genai

from docx import Document
from docx.shared import Pt
from io import BytesIO
import docx
from animations import display_cards

import os
from dotenv import load_dotenv

# Load API keys from .env
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

if not GEMINI_API_KEY:
    raise ValueError("❌ GEMINI_API_KEY not found in .env file")

# Configure Gemini API
genai.configure(api_key=GEMINI_API_KEY)


# -------------------- Helper / Cached functions --------------------
@st.cache_data
def load_data(file):
    return pd.read_csv(file)


@st.cache_data
def get_suggestions(student_name, marks_data, attendance_data):
    prompt = f"""
    Student Name: {student_name}
    Subject Marks: {marks_data}
    Attendance: {attendance_data}%

    As a teacher, provide personalized suggestions for this student to improve their performance (max 100 words and in bullet points):
    - Identify strengths and weaknesses based on subject marks
    - Appreciate for subjects the student performed well
    - Recommend subject-specific study strategies where the student is weak 
    - Address attendance issues if present
    - Suggestions must be specific to the student's performance
    - Suggest ways to maintain or boost motivation
    Max 4 bullet points.
    """
    model = genai.GenerativeModel("gemini-2.0-flash")
    response = model.generate_content(prompt)
    return response.text.strip() if response and hasattr(response, 'text') else "No suggestions generated."


@st.cache_data
def get_class_suggestions(subject_avgs):
    prompt = f"""
    Class Subject Averages: {subject_avgs}

    As a teacher, provide brief suggestions to improve overall class performance (max 50 words and in bullet points not more than 3):
    - Identify subjects where students are struggling
    - Recommend teaching strategies to improve these subjects
    - Suggest activities or resources to help students understand difficult concepts
    - Provide general tips to maintain or boost class motivation
    """
    model = genai.GenerativeModel("gemini-2.0-flash")
    response = model.generate_content(prompt)
    return response.text.strip() if response and hasattr(response, "text") else "No suggestions generated."


def calculate_performance(marks):
    return sum(marks) / len(marks)


def plot_performance(subjects, marks, title):
    fig, ax = plt.subplots(figsize=(10, 6))
    sns.barplot(x=subjects, y=marks, palette="coolwarm", ax=ax)
    ax.set_title(title, fontsize=16)
    ax.set_ylim(0, 100)
    for p in ax.patches:
        ax.annotate(f'{p.get_height():.2f}', 
                    (p.get_x() + p.get_width() / 2., p.get_height()),
                    ha='center', va='center', 
                    xytext=(0, 9), 
                    textcoords='offset points', 
                    fontsize=12)
    ax.set_xlabel('Subjects', fontsize=14)
    ax.set_ylabel('Marks', fontsize=14)
    sns.despine(fig)
    return fig


def analyze_subject_performance(df, subjects):
    weak_subjects = []
    strong_subjects = []
    avg_marks = df[subjects].mean()
    for subject in subjects:
        if avg_marks[subject] < 60:
            weak_subjects.append(subject)
        else:
            strong_subjects.append(subject)
    return weak_subjects, strong_subjects, avg_marks


@st.cache_data
def get_subject_suggestions(subject):
    prompt = f"""
    The class is struggling in {subject}. Provide brief strategies to help students improve in this subject (50 words max) in max 3 bullet points.
    """
    model = genai.GenerativeModel("gemini-2.0-flash")
    response = model.generate_content(prompt)
    return response.text.strip() if response and hasattr(response, "text") else "No suggestions generated."


def attendance_insights(df):
    avg_attendance = df['Attendance'].mean()
    min_attendance = df['Attendance'].min()
    max_attendance = df['Attendance'].max()
    avg_marks = df[[col for col in df.columns if col not in ['Roll No', 'Name', 'Attendance']]].mean(axis=1)
    correlation = df['Attendance'].corr(avg_marks)
    if correlation > 0.5:
        attendance_impact = "Low attendance is significantly impacting performance. Ensure regular attendance."
    elif correlation > 0:
        attendance_impact = "Attendance is moderately impacting performance. Try to attend more regularly."
    else:
        attendance_impact = "Attendance is not a major issue for performance. Focus on study habits."
    lowest_attendance_student = df[df['Attendance'] == min_attendance]['Name'].values[0]
    highest_attendance_student = df[df['Attendance'] == max_attendance]['Name'].values[0]
    insights = f"""
    - Average Attendance: {avg_attendance:.2f}%
    - Lowest Attendance: {min_attendance}% (Student: {lowest_attendance_student})
    - Highest Attendance: {max_attendance}% (Student: {highest_attendance_student})
    - Insights: {attendance_impact}
    """
    return insights


def save_insights_to_docx(title, insights, charts):
    doc = Document()
    doc.add_heading(title, level=1)
    for insight in insights.split('\n'):
        if insight.strip():
            p = doc.add_paragraph(insight.strip(), style='BodyText')
            for run in p.runs:
                run.font.size = Pt(12)
    for chart in charts:
        image_stream = BytesIO()
        chart.savefig(image_stream, format='png')
        image_stream.seek(0)
        doc.add_picture(image_stream, width=docx.shared.Inches(6))
    return doc


# -------------------- New: query_gemini function (fixes NameError) --------------------
def query_gemini(question: str, df: pd.DataFrame) -> str:
    """
    Build a concise dataset summary and ask Gemini the question.
    We intentionally include only a short summary + first 10 rows to keep prompts small.
    """
    # Basic metadata
    n_rows = len(df)
    columns = df.columns.tolist()

    # Numeric summary (if present)
    numeric_df = df.select_dtypes(include=['number'])
    if not numeric_df.empty:
        numeric_summary = numeric_df.describe().round(2).to_string()
    else:
        numeric_summary = "No numeric columns."

    # Head (limit to first 10 rows)
    head = df.head(10).to_string(index=False)

    prompt = f"""
You are a helpful assistant that answers questions about the dataset provided.

Dataset summary:
- Rows: {n_rows}
- Columns: {columns}

Numeric summary (statistics for numeric columns):
{numeric_summary}

First up to 10 rows of the table:
{head}

Question:
{question}

Answer concisely, cite column names or row examples where relevant. If the question cannot be answered from the dataset, say you don't have enough information.
"""
    try:
        model = genai.GenerativeModel("gemini-2.0-flash")
        response = model.generate_content(prompt)
        return response.text.strip() if response and hasattr(response, "text") else "No answer generated."
    except Exception as e:
        return f"Error when calling Gemini: {e}"


# -------------------- Main analysis / Streamlit UI --------------------
def analysis():
    uploaded_file = st.file_uploader("Upload CSV file with student data", type="csv")
    analysis_type = st.sidebar.radio(
        "Choose Analysis Type:",
        ["Class Wide Performance Analysis", "Student Wise Performance Analysis", "Attendance Analysis", "Ask Questions To The Data"],
        horizontal=False
    )

    if uploaded_file is not None:
        df = load_data(uploaded_file)
        required_columns = ['Roll No', 'Name', 'Attendance']

        if not all(col in df.columns for col in required_columns):
            st.error("CSV file must contain 'Roll No', 'Name', and 'Attendance' columns.")
            return

        subjects = [col for col in df.columns if col not in required_columns]

        # -------------------- STUDENT-WISE PERFORMANCE --------------------
        if analysis_type == "Student Wise Performance Analysis":
            st.markdown("<h1 style='font-size:30px;font-family:Garamond,serif;'>Student-wise Analysis</h1>", unsafe_allow_html=True)
            student_names = df['Name'].unique()
            selected_student = st.selectbox("Select a student to analyze:", student_names)

            student_data = df[df['Name'] == selected_student].iloc[0]
            marks = {subject: student_data[subject] for subject in subjects}
            attendance = student_data['Attendance']
            overall_score = calculate_performance(list(marks.values()))

            st.markdown(f"<h1 style='font-size:30px;font-family:Garamond,serif;'>{selected_student}'s Performance</h1>", unsafe_allow_html=True)
            st.write(f"Average Score: {overall_score:.2f}/100")
            st.write(f"Attendance: {attendance}%")

            fig = plot_performance(subjects, list(marks.values()), f"{selected_student}'s Subject-wise Marks")
            st.pyplot(fig)

            categories = {
                'Excellent': 90,
                'Good': 80,
                'Needs Improvement': 60,
                'Concerning': 40,
                'Failed': 0
            }

            st.markdown("<h1 style='font-size:30px;font-family:Garamond,serif;'>Subject-wise Status</h1>", unsafe_allow_html=True)
            for subject, mark in marks.items():
                for cat, threshold in categories.items():
                    if mark >= threshold:
                        st.write(f"{subject}: {cat} ({mark}/100)")
                        break

            if attendance < 50:
                st.error("🚨 CRITICAL WARNING: Attendance is dangerously low. Immediate action is required.")
            elif attendance < 75:
                st.warning("⚠ Attendance is below 75%. This can significantly impact performance.")

            st.markdown("<h1 style='font-size:30px;font-family:Garamond,serif;'>Personalized Suggestions</h1>", unsafe_allow_html=True)
            suggestions = get_suggestions(selected_student, marks, attendance)
            st.write(suggestions)

            charts = [fig]
            doc = save_insights_to_docx(f"{selected_student}'s Performance Insights", suggestions, charts)
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button(
                label="Download Student Insights",
                data=buffer,
                file_name=f"{selected_student}_insights.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        # -------------------- CLASS-WIDE PERFORMANCE --------------------
        elif analysis_type == "Class Wide Performance Analysis":
            st.markdown("<h1 style='font-size:30px;font-family:Garamond,serif;'>Class-wide Analysis</h1>", unsafe_allow_html=True)

            class_avg = df[subjects].mean().mean()
            weak_subjects, strong_subjects, avg_marks = analyze_subject_performance(df, subjects)

            st.markdown("<h1 style='font-size:30px;font-family:Garamond,serif;'>Subjects Analysis</h1>", unsafe_allow_html=True)
            st.write("Subjects where students are performing well:")
            for subject in strong_subjects:
                st.write(f"- {subject}: {avg_marks[subject]:.2f}/100")

            st.write("Subjects where students are struggling:")
            for subject in weak_subjects:
                st.write(f"- {subject}: {avg_marks[subject]:.2f}/100")

            avg_marks_series = df[subjects].mean()
            highest_marks = df[subjects].max()
            lowest_marks = df[subjects].min()

            # display_cards assumed to be your UI animation helper
            try:
                display_cards("Class Subject Performance", avg_marks_series.mean(), highest_marks.max(), lowest_marks.min())
            except Exception:
                # if animations fail, ignore and continue
                pass

            if weak_subjects:
                selected_subject = st.selectbox("Select a weak subject to get improvement suggestions:", weak_subjects)
                if selected_subject:
                    st.write(f"*Suggestions to Improve Performance in {selected_subject}:*")
                    subject_suggestions = get_subject_suggestions(selected_subject)
                    st.write(subject_suggestions)
            else:
                st.write("No weak subjects detected.")

            st.markdown("<h1 style='font-size:30px;font-family:Garamond,serif;'>Overall Class Improvement Plan</h1>", unsafe_allow_html=True)
            class_suggestions = get_class_suggestions(avg_marks.to_dict())
            st.write(class_suggestions)

            class_doc = Document()
            class_doc.add_heading("Class-wide Performance Insights", level=1)
            class_doc.add_heading("Subjects Analysis", level=2)

            class_doc.add_heading("Subjects where students are performing well:", level=3)
            for subject in strong_subjects:
                p = class_doc.add_paragraph(f"- {subject}: {avg_marks[subject]:.2f}/100", style='BodyText')
                for run in p.runs:
                    run.font.size = Pt(12)

            class_doc.add_heading("Subjects where students are struggling:", level=3)
            for subject in weak_subjects:
                p = class_doc.add_paragraph(f"- {subject}: {avg_marks[subject]:.2f}/100", style='BodyText')
                for run in p.runs:
                    run.font.size = Pt(12)

            class_doc.add_heading("Overall Class Improvement Plan", level=2)
            p = class_doc.add_paragraph(class_suggestions, style='BodyText')
            for run in p.runs:
                run.font.size = Pt(12)

            buffer = BytesIO()
            class_doc.save(buffer)
            buffer.seek(0)
            st.download_button(
                label="Download Class Insights",
                data=buffer,
                file_name="class_insights.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        # -------------------- ATTENDANCE ANALYSIS --------------------
        elif analysis_type == "Attendance Analysis":
            st.markdown("<h1 style='font-size:30px;font-family:Garamond,serif;'>Attendance Analysis</h1>", unsafe_allow_html=True)
            attendance_report = attendance_insights(df)
            st.write(attendance_report)

            fig, ax = plt.subplots(figsize=(8, 5))
            sns.histplot(df['Attendance'], bins=10, kde=True, ax=ax)
            ax.set_title("Attendance Distribution", fontsize=16)
            ax.set_xlabel("Attendance (%)")
            ax.set_ylabel("Number of Students")
            st.pyplot(fig)

            charts = [fig]
            doc = save_insights_to_docx("Attendance Analysis", attendance_report, charts)
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button(
                label="Download Attendance Insights",
                data=buffer,
                file_name="attendance_insights.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        # -------------------- ASK QUESTIONS TO THE DATA --------------------
        elif analysis_type == "Ask Questions To The Data":
            st.markdown("<h1 style='font-size:30px;font-family:Garamond,serif;'>Ask Questions To The Data</h1>", unsafe_allow_html=True)
            st.write("You can ask questions about the dataset (examples: 'Which student has the lowest attendance?', 'Average marks in Maths', 'How many students scored below 40 in Science?').")

            question = st.text_area("Ask a question about the dataset:")
            if st.button("Get Answer"):
                if not question.strip():
                    st.warning("Please type a question before clicking 'Get Answer'.")
                else:
                    with st.spinner("Querying Gemini..."):
                        answer = query_gemini(question, df)
                    # display answer with a success box
                    st.success("Answer from Gemini:")
                    st.write(answer)

    else:
        st.info("Please upload a CSV file with the required columns: Roll No, Name, Attendance, and at least one subject column.")


if __name__ == "__main__":
    analysis()
